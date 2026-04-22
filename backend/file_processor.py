import base64
import io
import PyPDF2
import docx

def process_file(file_bytes: bytes, filename: str) -> dict:
    ext = filename.split('.')[-1].lower() if '.' in filename else ''
    
    result = {
        "text": None,
        "image_base64": None,
        "file_type": ext
    }
    
    if ext == 'pdf':
        try:
            reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            text = ""
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            
            if text.strip():
                result["text"] = text
            else:
                result["text"] = ""
        except Exception as e:
            result["text"] = f"Error reading PDF: {e}"
            
    elif ext == 'docx':
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            text = "\n".join([para.text for para in doc.paragraphs])
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    text += " | ".join([cell.text for cell in row.cells]) + "\n"
            result["text"] = text
        except Exception as e:
            result["text"] = f"Error reading DOCX: {e}"
            
    elif ext in ['txt', 'csv']:
        try:
            result["text"] = file_bytes.decode('utf-8')
        except UnicodeDecodeError:
            # Fallback to general decoding
            result["text"] = file_bytes.decode('latin-1', errors='ignore')
            
    elif ext in ['png', 'jpg', 'jpeg']:
        b64_str = base64.b64encode(file_bytes).decode('utf-8')
        result["image_base64"] = b64_str
    
    else:
        result["text"] = f"Unsupported file type: {ext}"
        
    return result
